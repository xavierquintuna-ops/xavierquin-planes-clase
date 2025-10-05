# -*- coding: utf-8 -*-
"""
app.py - Generador de Plan de Clase (Versión final)
Incluye:
 - Corrección de errores de codificación UTF-8
 - Generación en inglés si la asignatura es "Inglés / English"
 - Contador estimado de tokens por planificación + historial y gráfica
 - Estilo visual personalizado (CSS)
 - Banner superior animado con frases motivadoras aleatorias
 - Descarga a Word y Excel
 - Interfaz organizada con expanders y mensajes claros
"""

import streamlit as st
from io import BytesIO
from docx import Document
import time
import unicodedata
import datetime
import random
from typing import List, Dict, Any
import pandas as pd
import matplotlib.pyplot as plt

# -------------------------
# Dependencia para Gemini (Google Generative AI)
# Asegúrate de tener instalada la librería google-genai y configurada.
# -------------------------
from google import genai
from google.genai.errors import APIError

# -------------------------
# CONFIG: coloca tu API KEY aqui (no la expongas en frontend)
# -------------------------
GEMINI_API_KEY = "AIzaSyC0FOYvSIwW2WEePc4ks_dB6WdHyVBvmy0"  # reemplaza por tu clave real antes de desplegar
MODEL_NAME = "gemini-2.5-flash"
MAX_TOKENS = 2800
TEMPERATURE = 0.3

# -------------------------
# Página y estilo general
# -------------------------
st.set_page_config(page_title="Planificador Educativo", page_icon="📘", layout="wide")

# Frases motivadoras aleatorias
frases_docentes = [
    "Educar es sembrar esperanza 🌱",
    "El mejor maestro enseña con el corazón ❤️",
    "Compartir conocimiento es dejar huella ✨",
    "Cada clase es una oportunidad para transformar vidas 🌍",
    "La educación es el arma más poderosa para cambiar el mundo 🌟",
    "Un docente inspira más allá de las palabras 💡",
    "La enseñanza que deja huella va de corazón a corazón 💖"
]
frase_motivadora = random.choice(frases_docentes)

# CSS personalizado + banner animado
custom_css = f"""
<style>
/* Fondo y tipografía */
.stApp {{
    background: linear-gradient(135deg, #eaf3ff, #ffffff);
    color: #222;
    font-family: 'Segoe UI', Roboto, 'Helvetica Neue', Arial, sans-serif;
    padding-top: 10px;
}}

/* Banner animado */
.banner {{
    position: relative;
    height: 220px;
    background: linear-gradient(135deg, #1a73e8, #4285f4);
    overflow: hidden;
    border-radius: 12px;
    margin-bottom: 20px;
    box-shadow: 0 6px 22px rgba(0,0,0,0.15);
    text-align: center;
    padding-top: 30px;
}}

/* Texto del banner */
.banner h1 {{
    color: white;
    font-size: 30px;
    font-weight: 700;
    margin: 0;
    z-index: 2;
    position: relative;
    letter-spacing: 0.2px;
}}
.banner h2 {{
    color: #f1f1f1;
    font-size: 16px;
    font-weight: 400;
    margin-top: 8px;
    z-index: 2;
    position: relative;
    font-style: italic;
}}

/* Onda animada */
.wave {{
    position: absolute;
    bottom: 0;
    left: 0;
    width: 200%;
    height: 100%;
    background-repeat: repeat-x;
    background-size: 50% 100%;
    opacity: 0.55;
    animation: move 12s linear infinite;
    z-index: 1;
}}
.wave1 {{
    background-image: radial-gradient(circle at 50% 40%, rgba(255,255,255,0.35) 15%, transparent 60%);
    height: 100%;
}}
.wave2 {{
    background-image: radial-gradient(circle at 50% 50%, rgba(255,255,255,0.18) 12%, transparent 60%);
    height: 120%;
    animation-duration: 18s;
}}

@keyframes move {{
    0% {{ transform: translateX(0); }}
    100% {{ transform: translateX(-25%); }}
}}

/* Inputs y botones */
.stTextInput > div > div > input, .stTextArea textarea {{
    border: 1px solid #1a73e8 !important;
    border-radius: 8px !important;
    background-color: #fbfeff !important;
    padding: 8px !important;
}}
button, .stButton>button {{
    background-color: #1a73e8 !important;
    color: white !important;
    border-radius: 10px !important;
    padding: 6px 12px !important;
    font-size: 14px !important;
}
button:hover, .stButton>button:hover {{
    background-color: #1557b0 !important;
}}

/* Mensajes */
.stAlert {{
    border-radius: 8px !important;
}}
</style>

<div class="banner">
    <h1>📘 XAVIERQUIN - Planificación de Clases</h1>
    <h2>{frase_motivadora}</h2>
    <div class="wave wave1"></div>
    <div class="wave wave2"></div>
</div>
"""
st.markdown(custom_css, unsafe_allow_html=True)

# -------------------------
# Inicialización de session_state
# -------------------------
defaults = {
    "asignatura": "",
    "grado": "",
    "edad": 12,
    "tema_insercion": "",
    "destrezas": [],
    "plan_text": None,
    "doc_bytes": None,
    "excel_bytes": None,
    "last_error": "",
    "tokens_usados": 0,
    "historial_tokens": [],
    "planes_generados": 0
}
for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

# -------------------------
# Utilidades
# -------------------------
def normalize_text(s: str) -> str:
    if s is None:
        return ""
    return unicodedata.normalize("NFKC", str(s)).strip()

def contar_tokens_estimado(texto: str) -> int:
    """
    Estimación simple de tokens: 1 token ≈ 4 caracteres (esta es una aproximación).
    Es útil para control de consumo aproximado.
    """
    if not texto:
        return 0
    return max(1, len(texto) // 4)

def create_docx_from_text(plan_text: str) -> BytesIO:
    doc = Document()
    doc.add_heading("Plan de Clase", level=1)
    for line in plan_text.split("\n"):
        if line.strip():
            doc.add_paragraph(line)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def create_excel_from_plan(destrezas: List[Dict[str,str]], plan_text: str) -> BytesIO:
    rows = []
    for d in destrezas:
        rows.append({
            "DESTREZA": d.get("destreza", ""),
            "INDICADOR": d.get("indicador", ""),
            "TEMA": d.get("tema_estudio", ""),
            "PLAN DE CLASE": plan_text
        })
    df = pd.DataFrame(rows)
    buf = BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Planificacion")
    buf.seek(0)
    return buf

# -------------------------
# Llamada al modelo Gemini
# -------------------------
def call_model(prompt_text: str) -> str:
    if not GEMINI_API_KEY or GEMINI_API_KEY == "TU_API_KEY_AQUI":
        raise RuntimeError("La clave API de Gemini no está configurada en el código. Reemplaza GEMINI_API_KEY.")
    try:
        client = genai.Client(api_key=GEMINI_API_KEY)
        config = genai.types.GenerateContentConfig(
            temperature=TEMPERATURE,
            max_output_tokens=MAX_TOKENS,
        )
        response = client.models.generate_content(
            model=MODEL_NAME,
            contents=[{"role": "user", "parts": [{"text": prompt_text}]}],
            config=config,
        )
        # response.text devuelve el texto generado
        return response.text
    except APIError as e:
        # Mensaje claro para el usuario (no exponer detalles sensibles)
        st.error(f"Error con la API de Gemini: {e}")
        raise
    except Exception as e:
        st.error(f"Error inesperado: {e}")
        raise

# -------------------------
# Construcción del prompt (soporte Inglés)
# -------------------------
def build_prompt(asignatura: str, grado: str, edad: Any, tema_insercion: str, destrezas_list: List[Dict[str,str]]) -> str:
    is_english = asignatura.strip().lower() in ["ingles", "inglés", "english"]

    if is_english:
        instructions = (
            "You are an expert in curriculum design and lesson planning. Generate a LESSON PLAN in clear U.S. ENGLISH. "
            "Provide a structured and practical lesson plan including anticipation, construction, consolidation, resources and evaluation.\n\n"
            f"Subject: {asignatura}\n"
            f"Grade: {grado}\n"
            f"Age: {edad}\n"
            f"Transversal Topic: {tema_insercion}\n\n"
            "### SKILLS AND INDICATORS\n"
        )
        for d in destrezas_list:
            instructions += f"- Skill: {d.get('destreza','')} | Indicator: {d.get('indicador','')}\n"

        instructions += (
            "\n### ANTICIPATION\n- Activities that activate prior knowledge (brief, action-oriented).\n\n"
            "### CONSTRUCTION\n- At least 6 sequenced activities in progressive order.\n\n"
            "### CONSOLIDATION\n- Activities to apply and reinforce knowledge.\n\n"
            "### RESOURCES\n- List physical and technological resources.\n\n"
            "### EVALUATION\n- Clear evaluation activities aligned to indicators.\n\n"
        )
    else:
        instructions = (
            "Eres un experto en diseño curricular y planificación educativa. Genera un PLAN DE CLASE en ESPAÑOL, "
            "estructurado y práctico. Incluye anticipación, construcción, consolidación, recursos y evaluación.\n\n"
            f"Asignatura: {asignatura}\n"
            f"Grado: {grado}\n"
            f"Edad: {edad}\n"
            f"Tema de Inserción: {tema_insercion}\n\n"
            "### DESTREZAS E INDICADORES\n"
        )
        for d in destrezas_list:
            instructions += f"- Destreza: {d.get('destreza','')} | Indicador: {d.get('indicador','')}\n"

        instructions += (
            "\n### ANTICIPACIÓN\n- Actividades que activen conocimientos previos (breves, con verbos en infinitivo si es posible).\n\n"
            "### CONSTRUCCIÓN\n- Al menos 6 actividades en secuencia pedagógica.\n\n"
            "### CONSOLIDACIÓN\n- Actividades para aplicar y reforzar aprendizajes.\n\n"
            "### RECURSOS\n- Indicar recursos físicos y tecnológicos.\n\n"
            "### EVALUACIÓN\n- Actividades de evaluación alineadas con los indicadores.\n\n"
        )

    return instructions

# -------------------------
# Interfaz - Entrada de datos
# -------------------------
with st.expander("📋 Ingresar datos básicos", expanded=True):
    c1, c2 = st.columns(2)
    with c1:
        st.text_input("Asignatura", key="asignatura", value=st.session_state["asignatura"])
        st.text_input("Grado", key="grado", value=st.session_state["grado"])
    with c2:
        st.number_input("Edad de los estudiantes", min_value=3, max_value=99, key="edad", value=st.session_state["edad"])
        st.text_input("Tema de Inserción (actividad transversal)", key="tema_insercion", value=st.session_state["tema_insercion"])

st.markdown("---")
st.subheader("➕ Agregar destreza e indicador")

with st.form(key="form_add_destreza"):
    d = st.text_area("Destreza", key="form_destreza")
    i = st.text_area("Indicador de logro", key="form_indicador")
    t = st.text_input("Tema de estudio (opcional)", key="form_tema_estudio")
    submitted = st.form_submit_button("➕ Agregar destreza")
    if submitted:
        dd, ii, tt = normalize_text(d), normalize_text(i), normalize_text(t)
        if not dd or not ii:
            st.warning("⚠️ Completa la destreza y el indicador antes de agregar.")
        else:
            st.session_state["destrezas"].append({"destreza": dd, "indicador": ii, "tema_estudio": tt})
            st.success("✔️ Destreza agregada")
            st.rerun()

if st.session_state["destrezas"]:
    st.subheader("📌 Destrezas añadidas")
    st.table(st.session_state["destrezas"])

# -------------------------
# Generar plan - callback
# -------------------------
def generar_plan_callback():
    st.session_state["last_error"] = ""
    asig = normalize_text(st.session_state["asignatura"])
    grad = normalize_text(st.session_state["grado"])
    edad_val = st.session_state["edad"]
    tema = normalize_text(st.session_state["tema_insercion"])
    dests = st.session_state["destrezas"]

    if not asig or not grad or not dests:
        st.session_state["last_error"] = "⚠️ Faltan campos obligatorios (asignatura, grado o al menos 1 destreza)."
        return

    try:
        with st.spinner("⏳ Generando plan con Gemini..."):
            prompt = build_prompt(asig, grad, edad_val, tema, dests)
            respuesta = call_model(prompt)

        # Guardar respuesta y archivos
        st.session_state["plan_text"] = respuesta
        st.session_state["doc_bytes"] = create_docx_from_text(respuesta).getvalue()
        st.session_state["excel_bytes"] = create_excel_from_plan(dests, respuesta).getvalue()

        # Estimación de tokens (aprox.)
        tokens_estimados = contar_tokens_estimado(respuesta)
        st.session_state["tokens_usados"] += tokens_estimados
        st.session_state["historial_tokens"].append(tokens_estimados)
        st.session_state["planes_generados"] += 1

        st.success(f"✔️ Plan generado con éxito. Tokens estimados usados: {tokens_estimados}")

    except Exception as e:
        st.session_state["last_error"] = str(e)

st.button("📄 Generar Plan de Clase", on_click=generar_plan_callback)

if st.session_state.get("last_error"):
    st.error(st.session_state["last_error"])

# -------------------------
# Vista previa del plan
# -------------------------
if st.session_state.get("plan_text"):
    st.markdown("---")
    st.subheader("👀 Vista previa del Plan")
    # Mostrar con markdown (preserva saltos de línea)
    st.markdown(st.session_state["plan_text"])

# -------------------------
# Exportar a Word y Excel
# -------------------------
if st.session_state.get("doc_bytes"):
    ts = time.strftime("%Y%m%d_%H%M%S")
    st.download_button(
        "💾 Exportar a Word",
        data=st.session_state["doc_bytes"],
        file_name=f"plan_{ts}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

if st.session_state.get("excel_bytes"):
    ts = time.strftime("%Y%m%d_%H%M%S")
    st.download_button(
        "📊 Exportar a Excel",
        data=st.session_state["excel_bytes"],
        file_name=f"plan_{ts}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

# -------------------------
# Panel de control: consumo de tokens y métricas
# -------------------------
st.markdown("---")
st.subheader("📈 Consumo y métricas de la sesión")

col1, col2, col3 = st.columns([2,1,1])
with col1:
    st.info(f"🔢 Tokens totales estimados consumidos en esta sesión: **{st.session_state['tokens_usados']}**")
    st.caption(f"🧾 Planes generados en esta sesión: **{st.session_state['planes_generados']}**")
with col2:
    if st.session_state["historial_tokens"]:
        st.metric("Último consumo (tokens)", st.session_state["historial_tokens"][-1])
    else:
        st.metric("Último consumo (tokens)", 0)
with col3:
    # Estimación simple de cuántas planificaciones cabrían en 1 millón tokens (referencia)
    if st.session_state['tokens_usados'] > 0:
        restante_estimado = max(0, 1000000 - st.session_state['tokens_usados'])
    else:
        restante_estimado = 1000000
    st.caption("Referencia: tokens/mes (ejemplo 1,000,000)")

# Gráfica historial
if st.session_state["historial_tokens"]:
    fig, ax = plt.subplots(figsize=(6, 3))
    ax.plot(st.session_state["historial_tokens"], marker="o")
    ax.set_title("Tokens estimados por planificación")
    ax.set_xlabel("N° planificación (sesión)")
    ax.set_ylabel("Tokens estimados")
    ax.grid(axis='y', linestyle='--', alpha=0.4)
    st.pyplot(fig)

# -------------------------
# Reiniciar aplicación
# -------------------------
def reset_app():
    for k, v in defaults.items():
        st.session_state[k] = v
    st.rerun()

st.markdown("---")
st.button("🆕 Nuevo (reiniciar formulario)", on_click=reset_app)

# -------------------------
# Sugerencias y ayuda rápida
# -------------------------
st.markdown("---")
with st.expander("❓ Consejos para ahorrar tokens y buenas prácticas", expanded=False):
    st.write("""
    - Reduce la longitud de la `Asigantura` o `Destreza` si no es necesaria mucha descripción.
    - Acota `max_output_tokens` en la configuración si quieres respuestas más cortas.
    - Valida automáticamente los links y recursos que la IA proponga (la IA puede inventar URLs).
    - Guarda la API key en variables de entorno o en Secret Manager (no en el repositorio).
    - Si vas a usar esta app con 56 docentes y 20 planes cada uno (~1,120 planes),
      estima el consumo: cada plan puede usar entre 800 y 1,500 tokens según detalle.
    """)

st.caption("ℹ️ Código generado por el asistente. Recuerda reemplazar GEMINI_API_KEY por tu clave real y protegerla en producción.")
